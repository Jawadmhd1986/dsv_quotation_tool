from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os
import re
from datetime import datetime

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("form.html")

@app.route("/generate", methods=["POST"])
def generate():
    storage_type = request.form["storage_type"]
    volume = float(request.form["volume"])
    days = int(request.form["days"])
    include_wms = request.form["wms"] == "Yes"
    email = request.form.get("email", "")
    today_str = datetime.today().strftime("%d %b %Y")

    if "chemical" in storage_type.lower():
        template_path = "templates/Chemical VAS.docx"
    elif "open yard" in storage_type.lower():
        template_path = "templates/Open Yard VAS.docx"
    else:
        template_path = "templates/Standard VAS.docx"

    doc = Document(template_path)

    if storage_type == "AC":
        rate = 2.5
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Non-AC":
        rate = 2.0
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Open Shed":
        rate = 1.8
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Chemicals AC":
        rate = 3.5
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Chemicals Non-AC":
        rate = 2.7
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif "kizad" in storage_type.lower():
        rate = 125
        unit = "SQM"
        rate_unit = "SQM / YEAR"
        storage_fee = volume * days * (rate / 365)
    elif "mussafah" in storage_type.lower():
        rate = 160
        unit = "SQM"
        rate_unit = "SQM / YEAR"
        storage_fee = volume * days * (rate / 365)
    else:
        rate = 0
        storage_fee = 0
        unit = "CBM"
        rate_unit = "CBM / DAY"

    storage_fee = round(storage_fee, 2)
    months = max(1, days // 30)
    is_open_yard = "open yard" in storage_type.lower()
    wms_fee = 0 if is_open_yard or not include_wms else 1500 * months
    total_fee = round(storage_fee + wms_fee, 2)

    placeholders = {
        "{{STORAGE_TYPE}}": storage_type,
        "{{DAYS}}": str(days),
        "{{VOLUME}}": str(volume),
        "{{UNIT}}": unit,
        "{{WMS_STATUS}}": "" if is_open_yard else ("INCLUDED" if include_wms else "NOT INCLUDED"),
        "{{UNIT_RATE}}": f"{rate:.2f} AED / {rate_unit}",
        "{{STORAGE_FEE}}": f"{storage_fee:,.2f} AED",
        "{{WMS_FEE}}": f"{wms_fee:,.2f} AED",
        "{{TOTAL_FEE}}": f"{total_fee:,.2f} AED",
        "{{TODAY_DATE}}": today_str
    }

    def replace_placeholders(doc, mapping):
        for p in doc.paragraphs:
            for key, val in mapping.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in mapping.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)

    replace_placeholders(doc, placeholders)

    def delete_block(doc, start_tag, end_tag):
        inside = False
        to_delete = []
        for i, p in enumerate(doc.paragraphs):
            if start_tag in p.text:
                inside = True
                to_delete.append(i)
            elif end_tag in p.text:
                to_delete.append(i)
                inside = False
            elif inside:
                to_delete.append(i)
        for i in reversed(to_delete):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    if "open yard" in storage_type.lower():
        delete_block(doc, "[VAS_STANDARD]", "[/VAS_STANDARD]")
        delete_block(doc, "[VAS_CHEMICAL]", "[/VAS_CHEMICAL]")
    elif "chemical" in storage_type.lower():
        delete_block(doc, "[VAS_STANDARD]", "[/VAS_STANDARD]")
        delete_block(doc, "[VAS_OPENYARD]", "[/VAS_OPENYARD]")
    else:
        delete_block(doc, "[VAS_CHEMICAL]", "[/VAS_CHEMICAL]")
        delete_block(doc, "[VAS_OPENYARD]", "[/VAS_OPENYARD]")

    os.makedirs("generated", exist_ok=True)
    filename_prefix = email.split('@')[0] if email else "quotation"
    filename = f"Quotation_{filename_prefix}.docx"
    output_path = os.path.join("generated", filename)
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    message = data.get("message", "").lower().strip()

    def normalize(text):
        text = text.lower().strip()

    # Common chat language
        text = re.sub(r"\bu\b", "you", text)
        text = re.sub(r"\bur\b", "your", text)
        text = re.sub(r"\br\b", "are", text)
        text = re.sub(r"\bpls\b", "please", text)
        text = re.sub(r"\bthx\b", "thanks", text)
        text = re.sub(r"\binfo\b", "information", text)

    # Logistics & warehouse short forms
        text = re.sub(r"\bwh\b", "warehouse", text)
        text = re.sub(r"\bw\/h\b", "warehouse", text)
        text = re.sub(r"\binv\b", "inventory", text)
        text = re.sub(r"\btemp\b", "temperature", text)
        text = re.sub(r"\btemp zone\b", "temperature zone", text)
        text = re.sub(r"\bwms system\b", "wms", text)

    # Transportation & locations
        text = re.sub(r"\brak\b", "ras al khaimah", text)
        text = re.sub(r"\babudhabi\b", "abu dhabi", text)
        text = re.sub(r"\babudhabi\b", "abu dhabi", text)
        text = re.sub(r"\bdxb\b", "dubai", text)

    # Industry abbreviations
        text = re.sub(r"\bo&g\b", "oil and gas", text)
        text = re.sub(r"\bdg\b", "dangerous goods", text)
        text = re.sub(r"\bfmcg\b", "fast moving consumer goods", text)

    # Quotation & VAS
        text = re.sub(r"\bdoc\b", "documentation", text)
        text = re.sub(r"\bdocs\b", "documentation", text)
        text = re.sub(r"\bmsds\b", "material safety data sheet", text)
        text = re.sub(r"\bvas\b", "value added services", text)

    # E-commerce variations
        text = re.sub(r"\be[\s\-]?commerce\b", "ecommerce", text)
        text = re.sub(r"\bshop logistics\b", "ecommerce", text)

    # Logistics models
        text = re.sub(r"\b3\.5pl\b", "three and half pl", text)
        text = re.sub(r"\b2pl\b", "second party logistics", text)
        text = re.sub(r"\b3pl\b", "third party logistics", text)
        text = re.sub(r"\b4pl\b", "fourth party logistics", text)

    # Fleet & vehicle types
        text = re.sub(r"\breefer\b", "refrigerated truck", text)
        text = re.sub(r"\bchiller\b", "refrigerated truck", text)
        text = re.sub(r"\bcity truck\b", "small truck", text)
        text = re.sub(r"\bev truck\b", "electric truck", text)

    # Fire system
        text = re.sub(r"\bfm200\b", "fm 200", text)

    # Misc business terms
        text = re.sub(r"\bkitting\b", "kitting and assembly", text)
        text = re.sub(r"\btagging\b", "labeling", text)
        text = re.sub(r"\basset tagging\b", "asset labeling", text)
        text = re.sub(r"\btransit store\b", "transit warehouse", text)
        text = re.sub(r"\basset mgmt\b", "asset management", text)
        text = re.sub(r"\bmidday break\b", "summer break", text)

    # Strip non-alphanumeric except spaces
        text = re.sub(r"[^a-z0-9\s]", "", text)

        return text

    message = normalize(message)

    def match(patterns):
        return any(re.search(p, message) for p in patterns)
        
# --- Containers (All Types + Flexible Unit Recognition) ---
    if match([r"healthcare|medical storage|pharma warehouse|pharma|pharmaceutical storage"]):
        return jsonify({"reply": "DSV serves healthcare clients via temperature-controlled, GDP-compliant storage at Abu Dhabi Airport Freezone and Mussafah."})

    if match([r"\bcontainers\b", r"tell me about containers", r"container types", r"types of containers", r"container sizes", r"container dimensions"]):
        return jsonify({"reply": "Here are the main container types and their specifications:\n\n📦 **20ft Container**:\n- Length: 6.1m, Width: 2.44m, Height: 2.59m\n- Payload: ~28,000 kg\n- Capacity: ~33 CBM\n\n📦 **40ft Container**:\n- Length: 12.2m, Width: 2.44m, Height: 2.59m\n- Payload: ~30,400 kg\n- Capacity: ~67 CBM\n\n⬆️ **40ft High Cube**:\n- Same as 40ft but height = 2.9m\n- Ideal for voluminous goods\n\n❄️ **Reefer Container (20ft & 40ft)**:\n- Insulated, temperature-controlled (+2°C to –25°C)\n- Used for food, pharma, perishables\n\n🏗 **Open Top Container**:\n- No roof, allows crane loading\n- For tall cargo (e.g. machinery, steel)\n\n🪜 **Flat Rack Container**:\n- No sides or roof\n- Used for oversized loads like vehicles or transformers\n\n📦 **SME Containers**:\n- Custom modular containers used in the UAE for small-scale import/export or temporary storage by SMEs\n\nLet me know if you'd like help choosing the right container for your cargo!"})

# Specific container types with ft/feet/foot flexibility
    if match([r"20\s*(ft|feet|foot)\s*container", r"\btwenty\s*(ft|feet|foot)?\s*container"]):
        return jsonify({"reply": "A 20ft container is 6.1m long × 2.44m wide × 2.59m high, capacity ~33 CBM, and payload up to 28,000 kg. Ideal for compact or heavy cargo."})

    if match([r"40\s*(ft|feet|foot)\s*container", r"\bforty\s*(ft|feet|foot)?\s*container"]):
        return jsonify({"reply": "A 40ft container is 12.2m long × 2.44m wide × 2.59m high, capacity ~67 CBM, and payload up to 30,400 kg. Suitable for palletized or bulk shipments."})

    if match([r"\bhigh cube\b.*container", r"40\s*(ft|feet|foot)\s*high cube", r"high cube container"]):
        return jsonify({"reply": "A 40ft High Cube container is 2.9m tall, 1 foot taller than standard containers. Ideal for bulky or voluminous cargo."})

    if match([r"reefer", r"refrigerated container", r"chiller container"]):
        return jsonify({"reply": "Reefer containers are temperature-controlled (+2°C to –25°C), ideal for perishables like food and pharmaceuticals. Available in 20ft and 40ft sizes."})

    if match([r"open top container", r"open roof", r"no roof container"]):
        return jsonify({"reply": "Open Top containers are used for tall or top-loaded cargo like steel coils, pipes, or machinery. They allow crane access from above."})

    if match([r"flat rack", r"no sides container", r"flat rack container"]):
        return jsonify({"reply": "Flat Rack containers have no sides or roof, perfect for oversized cargo such as vehicles, generators, or heavy equipment."})

    if match([r"\bsme\b", r"sme container", r"what is sme", r"sme size", r"sme container size"]):
        return jsonify({"reply": "In logistics, **SME** usually refers to Small and Medium Enterprises, but in UAE context, 'SME container' can also mean modular containers customized for SME use — often used for short-term cargo storage or small-scale import/export."})

    # --- Pallet Types, Sizes, and Positions ---
    if match([
    r"\bpallets\b", r"pallet types", r"types of pallets", r"pallet size", r"pallet sizes", r"pallet dimensions", 
    r"standard pallet", r"euro pallet", r"pallet specs", r"tell me about pallets", 
    r"what.*pallet.*used", r"pallet info", r"pallet.*per bay"]):
        return jsonify({"reply":
        "DSV uses two main pallet types in its 21K warehouse:\n\n"
        "🟦 **Standard Pallet**:\n- Size: 1.2m × 1.0m\n- Load capacity: ~1,000 kg\n- Fits **14 pallets per bay**\n\n"
        "🟨 **Euro Pallet**:\n- Size: 1.2m × 0.8m\n- Load capacity: ~800 kg\n- Fits **21 pallets per bay**\n\n"
        "Pallets are used for racking, picking, and transport. DSV also offers VAS like pallet loading, shrink wrapping, labeling, and stretch film wrapping for safe handling."})

# --- DSV UAE Divisions & Locations Overview ---
    if match([
    r"dsv.*uae", r"dsv.*locations.*uae", r"dsv.*in.*uae", r"dsv.*emirates", 
    r"dsv.*facilities.*uae", r"dsv.*sites.*uae", r"dsv.*divisions.*uae"]):
        return jsonify({"reply": 
        "**DSV UAE** operates logistics hubs across multiple Emirates:\n\n"
        "🏢 **Abu Dhabi (Main Operations)**:\n"
        "- 21K Warehouse (21,000 sqm, 7 chambers)\n"
        "- M44 / M45 (sub-warehouses in Mussafah)\n"
        "- Al Markaz (Hameem – 12,000 sqm regional site)\n"
        "- KIZAD Open Yard (360,000 sqm heavy cargo yard)\n"
        "- Airport Freezone (GDP-certified for healthcare)\n\n"
        "🚚 **Other Emirates:**\n"
        "- Dubai: Air & Sea coordination via Jebel Ali & DWC\n"
        "- Sharjah: SME and courier clients\n"
        "- Northern Emirates: Long-haul fleet transport\n\n"
        "Each site is integrated via WMS, VAS, and project logistics teams."})

# --- DSV Divisions / Global Business Units ---
    if match([
    r"dsv.*divisions", r"dsv.*business units", r"what are dsv.*divisions", r"dsv.*structure", 
    r"divisions of dsv", r"dsv.*department", r"dsv.*groups", r"dsv.*teams"]):
        return jsonify({"reply": 
        "DSV operates under three main global divisions:\n\n"
        "1️⃣ **DSV Air & Sea** – Freight forwarding by air and ocean\n"
        "2️⃣ **DSV Road** – Land transportation and fleet\n"
        "3️⃣ **DSV Solutions** – Contract logistics, warehousing, 3PL/4PL\n\n"
        "In the UAE, all divisions are active and managed under DSV Solutions PJSC (Abu Dhabi HQ)."})

# --- DSV UAE Air & Sea Details ---
    if match([
    r"air sea division", r"air & sea division", r"dsv air sea uae", r"air sea abu dhabi", 
    r"air sea freight uae", r"air sea forwarding uae", r"air sea.*dsv", r"dsv.*air.*sea"]):
        return jsonify({"reply": 
        "**DSV Air & Sea UAE** provides end-to-end forwarding via:\n\n"
        "✈️ **Air Freight**:\n"
        "- Standard, express, and charter cargo\n"
        "- Connected to Abu Dhabi Airport Freezone for healthcare & pharma\n\n"
        "🚢 **Sea Freight**:\n"
        "- FCL / LCL via Jebel Ali, Khalifa Port, and Zayed Port\n"
        "- Import/export clearance & project forwarding\n\n"
        "Operations are fully integrated with warehousing, transport, and WMS visibility."})

# --- DSV UAE Contact Info Block ---
    if match([
    r"dsv.*contact.*uae", r"dsv.*abu dhabi.*contact", r"dsv.*phone.*uae", 
    r"how to reach dsv uae", r"dsv.*email.*uae", r"contact dsv.*uae"]):
        return jsonify({"reply": 
        "**DSV Abu Dhabi (Head Office):**\n"
        "📍 Mussafah M19, 21K Warehouse\n"
        "📞 +971 2 555 2900\n"
        "📧 uae.sales@dsv.com\n\n"
        "For storage or quotation requests, please fill out the form on this page."})

    # --- All Storage Rates at Once ---
    if match([
    r"\ball\b",r"all.*storage.*rates", r"complete.*storage.*rate", r"all.*rate", r"list.*storage.*fees",
    r"storage.*rate.*overview", r"summary.*storage.*rates",
    r"show.*all.*storage.*charges", r"storage.*rates.*all", r"rates for all storage"]):
        return jsonify({"reply": 
        "**Here are the current DSV Abu Dhabi storage rates:**\n\n"
        "**📦 Standard Storage:**\n"
        "- AC: 2.5 AED/CBM/day\n"
        "- Non-AC: 2.0 AED/CBM/day\n"
        "- Open Shed: 1.8 AED/CBM/day\n\n"
        "**🧪 Chemical Storage:**\n"
        "- Chemical AC: 3.5 AED/CBM/day\n"
        "- Chemical Non-AC: 2.7 AED/CBM/day\n\n"
        "**🏗 Open Yard Storage:**\n"
        "- KIZAD: 125 AED/SQM/year\n"
        "- Mussafah: 160 AED/SQM/year\n\n"
        "*WMS fee applies to indoor storage unless excluded. For full quotation, fill out the form.*"})

    # --- Storage Rate Initial Question ---
    if match([r"storage rate[s]?$", r"\brates\b", r"storage", r"storage cost", r"how much.*storage", r"quotation.*storage only"]):
        return jsonify({"reply": "Which type of storage are you asking about? Standard, Chemicals, or Open Yard?"})

# --- Standard Storage Follow-up ---
    if match([r"^standard$", r"standard storage"]):
        return jsonify({"reply": "Do you mean Standard AC, Standard Non-AC, or Open Shed?"})

    if match([r"standard ac", r"ac standard", r"standard ac storage"]):
        return jsonify({"reply": "Standard AC storage is 2.5 AED/CBM/day. Standard VAS applies."})

    if match([r"standard non ac", r"non ac standard", r"standard non ac storage"]):
        return jsonify({"reply": "Standard Non-AC storage is 2.0 AED/CBM/day. Standard VAS applies."})

    if match([r"open shed", r"standard open shed", r"open shed storage rate"]):
        return jsonify({"reply": "Open Shed storage is 1.8 AED/CBM/day. Standard VAS applies."})

# --- Chemical Storage Follow-up ---
    if match([r"^chemical$", r"chemical storage only"]):
        return jsonify({"reply": "Do you mean Chemical AC or Chemical Non-AC?"})

    if match([r"chemical ac", r"ac chemical", r"chemical ac storage", r"chemical ac storage rate"]):
        return jsonify({"reply": "Chemical AC storage is 3.5 AED/CBM/day. Chemical VAS applies."})

    if match([r"chemical non ac", r"non ac chemical", r"chemical non ac storage", r"chemical non ac rate"]):
        return jsonify({"reply": "Chemical Non-AC storage is 2.7 AED/CBM/day. Chemical VAS applies."})

# --- Open Yard Storage ---
    if match([r"^open yard$", r"open yard storage", r"open yard rate", r"open yard storage rate"]):
        return jsonify({"reply": "Do you mean Open Yard in Mussafah or KIZAD?"})

    if match([r"open yard mussafah", r"mussafah open yard", r"rate.*mussafah open yard"]):
        return jsonify({"reply": "Open Yard Mussafah storage is **160 AED/SQM/year**. WMS is excluded. For availability, contact Antony Jeyaraj at antony.jeyaraj@dsv.com."})

    if match([r"open yard kizad", r"kizad open yard", r"rate.*kizad open yard"]):
        return jsonify({"reply": "Open Yard KIZAD storage is **125 AED/SQM/year**. WMS is excluded. For availability, contact Antony Jeyaraj at antony.jeyaraj@dsv.com."})
    if match([r"^kizad$", r"\bkizad\b"]):
        return jsonify({"reply": "Open Yard KIZAD storage is **125 AED/SQM/year**. WMS is excluded. For availability, contact Antony Jeyaraj at antony.jeyaraj@dsv.com."})
    if match([r"^mussafah$", r"\bmussafah\b"]):
        return jsonify({"reply": "Open Yard Mussafah storage is **160 AED/SQM/year**. WMS is excluded. For availability, contact Antony Jeyaraj at antony.jeyaraj@dsv.com."})

    # --- vas Rate ---
    if match([
    r"standard vas", r"standard", r"standard value added services", r"normal vas", r"normal value added services",
    r"handling charges", r"pallet charges", r"vas for ac", r"value added services for ac",
    r"vas for non ac", r"value added services for non ac",
    r"vas for open shed", r"value added services for open shed"]):
        return jsonify({"reply": "Standard VAS includes:\n- In/Out Handling: 20 AED/CBM\n- Pallet Loading: 12 AED/pallet\n- Documentation: 125 AED/set\n- Packing with pallet: 85 AED/CBM\n- Inventory Count: 3,000 AED/event\n- Case Picking: 2.5 AED/carton\n- Sticker Labeling: 1.5 AED/label\n- Shrink Wrapping: 6 AED/pallet\n- VNA Usage: 2.5 AED/pallet"})

    if match([
    r"chemical vas", r"chemical value added services",
    r"vas for chemical", r"value added services for chemical",
    r"hazmat vas", r"hazmat value added services",
    r"dangerous goods vas", r"dangerous goods value added services"]):
        return jsonify({"reply": "Chemical VAS includes:\n- Handling (Palletized): 20 AED/CBM\n- Handling (Loose): 25 AED/CBM\n- Documentation: 150 AED/set\n- Packing with pallet: 85 AED/CBM\n- Inventory Count: 3,000 AED/event\n- Inner Bag Picking: 3.5 AED/bag\n- Sticker Labeling: 1.5 AED/label\n- Shrink Wrapping: 6 AED/pallet"})

    if match([
    r"open yard vas", r"open yard", r"open yard value added services", r"yard equipment",
    r"forklift rate", r"crane rate", r"container lifting", r"yard charges"]):
        return jsonify({"reply": "Open Yard VAS includes:\n- Forklift (3T–7T): 90 AED/hr\n- Forklift (10T): 200 AED/hr\n- Forklift (15T): 320 AED/hr\n- Mobile Crane (50T): 250 AED/hr\n- Mobile Crane (80T): 450 AED/hr\n- Container Lifting: 250 AED/lift\n- Container Stripping (20ft): 1,200 AED/hr"})

    # --- Storage Rate Matching ---
    if match([r"open yard.*mussafah"]):
        return jsonify({"reply": "Open Yard Mussafah storage is 160 AED/SQM/year. WMS is excluded. VAS includes forklifts, cranes, and container lifts."})
    if match([r"open yard.*kizad"]):
        return jsonify({"reply": "Open Yard KIZAD storage is 125 AED/SQM/year. WMS excluded. VAS includes forklift 90–320 AED/hr, crane 250–450 AED/hr."})
   
    # --- 21K Warehouse  ---
    if match([r"rack height|rack levels|pallets per bay|racking"]):
        return jsonify({"reply": "21K warehouse racks are 12m tall with 6 pallet levels. Each bay holds 14 Standard pallets or 21 Euro pallets."})
    if match([r"\b21k\b", r"tell me about 21k", r"what is 21k", r"21k warehouse", r"21k dsv", r"main warehouse", r"mussafah.*21k"]):
        return jsonify({"reply": "21K is DSV’s main warehouse in Mussafah, Abu Dhabi. It is 21,000 sqm with a clear height of 15 meters. The facility features:\n- 3 rack types: Selective, VNA, and Drive-in\n- Rack height: 12m with 6 pallet levels\n- Aisle widths: Selective (2.95–3.3m), VNA (1.95m), Drive-in (2.0m)\n- 7 chambers used by clients like ADNOC, ZARA, PSN, and Civil Defense\n- Fully equipped with fire systems, access control, and RMS for document storage."})
    if match([r"\bgdsp\b", r"what is gdsp", r"gdsp certified", r"gdsp warehouse", r"gdsp compliance"]):
        return jsonify({"reply": "GDSP stands for Good Distribution and Storage Practices. It ensures that warehouse operations comply with global standards for the safe handling, storage, and distribution of goods, especially pharmaceuticals and sensitive materials. DSV’s warehouses in Abu Dhabi are GDSP certified."})
    if match([r"\biso\b", r"what iso", r"iso certified", r"tell me about iso", r"dsv iso", r"which iso standards"]):
        return jsonify({"reply": "DSV facilities in Abu Dhabi are certified with multiple ISO standards:\n- **ISO 9001**: Quality Management\n- **ISO 14001**: Environmental Management\n- **ISO 45001**: Occupational Health & Safety\nThese certifications ensure that DSV operates to the highest international standards in safety, service quality, and environmental responsibility."})
    if match([r"\bgdp\b", r"what is gdp", r"gdp warehouse", r"gdp compliant", r"gdp certified"]):
        return jsonify({"reply": "GDP stands for **Good Distribution Practice**, a quality standard for warehouse and transport operations of pharmaceutical products. DSV’s healthcare storage facilities in Abu Dhabi, including the Airport Freezone warehouse, are GDP-compliant, ensuring cold chain integrity, traceability, and regulatory compliance."})
    if match([r"cold chain", r"what.*cold chain", r"cold storage", r"temperature zones", r"what.*chains.*temperature", r"freezer room", r"cold room", r"ambient storage"]):
        return jsonify({"reply": "DSV offers full temperature-controlled logistics including:\n\n🟢 **Ambient Storage**: +18°C to +25°C (for general FMCG, electronics, and dry goods)\n🔵 **Cold Room**: +2°C to +8°C (for pharmaceuticals, healthcare, and food products)\n🔴 **Freezer Room**: –22°C (for frozen goods and sensitive biological materials)\n\nOur warehouses in Abu Dhabi are equipped with temperature monitoring, backup power, and GDP-compliant systems to maintain cold chain integrity."})
    if match([r"\brms\b", r"record management system", r"document storage", r"storage of files", r"paper storage"]):
        return jsonify({"reply": 
        "RMS (Record Management System) at DSV is located inside the 21K warehouse in Mussafah. It is used to store and manage physical documents, archives, and secure records for clients like Civil Defense.\n\n"
        "The RMS area is equipped with an **FM200 fire suppression system** for safe document protection. Note: RMS is not used for storing Return Material."})
    if match([r"asset management", r"what is asset management", r"tracking of assets", r"rfid.*asset"]):
        return jsonify({"reply": "DSV offers complete **Asset Management** solutions including:\n- Barcode or RFID tracking\n- Asset labeling\n- Storage and life-cycle monitoring\n- Secure location control\n\nIdeal for IT equipment, tools, calibration items, and government assets."})
    if match([r"quote.*asset", r"quotation.*asset management", r"what.*collect.*client.*asset", r"info.*for.*asset.*quotation"]):
        return jsonify({"reply":
        "To prepare an **Asset Management** quotation, collect the following from the client:\n"
        "1️⃣ Type of assets (IT, furniture, tools, etc.)\n"
        "2️⃣ Quantity and tagging type (barcode or RFID)\n"
        "3️⃣ Duration of storage or tracking\n"
        "4️⃣ Reporting/reporting system integration needs\n"
        "5️⃣ Any relocation, retrieval, or disposal cycles"})
    if match([r"asset labeling", r"asset labelling", r"label assets", r"tagging assets", r"rfid tagging", r"barcode tagging"]):
        return jsonify({"reply": "DSV provides **Asset Labeling** services using RFID or barcode tags. Labels include:\n- Unique ID numbers\n- Ownership info\n- Scannable codes for inventory and asset tracking\nThese are applied during intake or on-site at the client's request."})

    # --- 21K Warehouse Racking Info ---
    if match([
    r"\brack\b", r"\bracks\b", r"warehouse rack", r"warehouse racks", r"rack types",
    r"types of racks", r"racking system", r"rack system", r"racking layout", r"rack height",
    r"rack.*info", r"rack.*design", r"21k.*rack", r"rack.*21k", r"pallet levels"]):
        return jsonify({"reply":
        "The 21K warehouse in Mussafah uses 3 racking systems:\n\n"
        "🔷 **Selective Racking**:\n- Aisle width: 2.95m–3.3m\n- Standard access to all pallets\n\n"
        "🔷 **VNA (Very Narrow Aisle)**:\n- Aisle width: 1.95m\n- High-density storage with specialized forklifts\n\n"
        "🔷 **Drive-in Racking**:\n- Aisle width: 2.0m\n- Deep storage for uniform SKUs\n\n"
        "All racks are **12 meters tall** with **6 pallet levels plus ground**.\n"
        "Each bay holds:\n- **14 Standard pallets** (1.2m × 1.0m)\n- **21 Euro pallets** (1.2m × 0.8m)"})
    
    if match([
    r"pallet positions", r"how many.*pallet.*position", r"pallet slots", 
    r"positions per bay", r"rack.*pallet.*position", r"warehouse pallet capacity"]):
        return jsonify({"reply": 
        "Each rack bay in the 21K warehouse has:\n"
        "- **6 pallet levels** plus ground\n"
        "- Fits **14 Standard pallets** or **21 Euro pallets** per bay\n\n"
        "Across the facility, DSV offers thousands of pallet positions for ambient, VNA, and selective racking layouts. The exact total depends on rack type and client configuration."})
# --- Aisle Widths in 21K Warehouse ---
    if match([
    r"\baisle\b", r"aisle width", r"width of aisle", r"aisles", r"warehouse aisle", 
    r"vna aisle", r"how wide.*aisle", r"rack aisle width"]):
        return jsonify({"reply": 
        "Here are the aisle widths used in DSV’s 21K warehouse:\n\n"
        "🔹 **Selective Racking**: 2.95m – 3.3m\n"
        "🔹 **VNA (Very Narrow Aisle)**: 1.95m\n"
        "🔹 **Drive-in Racking**: 2.0m\n\n"
        "These widths are optimized for reach trucks, VNA machines, and efficient space utilization."})
   
# --- Warehouse Area / Size ---
    if match([
    r"\barea\b", r"warehouse area", r"warehouses area", r"warehouse size", r"warehouses size",
    r"how big.*warehouse", r"storage area",r"facilities",r"warehouses", r"warehouse total sqm", r"warehouse.*dimensions"]):
        return jsonify({"reply": 
        "DSV Abu Dhabi has approximately **44,000 sqm** of total warehouse space, distributed as follows:\n"
        "- **21K Warehouse (Mussafah)**: 21,000 sqm\n"
        "- **M44**: 5,760 sqm\n"
        "- **M45**: 5,000 sqm\n"
        "- **Al Markaz (Hameem)**: 12,000 sqm\n\n"
        "Additionally, we have **360,000 sqm** of open yard space, and a total logistics site of **481,000 sqm** including service roads and utilities."})

# --- Warehouse Space Availability ---
    if match([
    r"warehouse.*space.*available", r"do you have.*warehouse.*space", r"space in warehouse", 
    r"any warehouse space", r"warehouse availability", r"available.*storage", 
    r"available.*warehouse", r"wh space.*available", r"vacant.*warehouse"]):
        return jsonify({"reply": "For warehouse occupancy, please contact Biju Krishnan at **biju.krishnan@dsv.com**. He’ll assist with availability, allocation, and scheduling a site visit if needed."})
# --- Open Yard Space Availability ---
    if match([
    r"open yard.*occupancy", r"space.*open yard",r"space", r"open yard.*available", 
    r"do we have.*open yard", r"open yard availability", r"open yard.*space", 
    r"yard capacity", r"yard.*vacancy", r"any.*open yard.*space"]):
        return jsonify({"reply": "For open yard occupancy, please contact Antony Jeyaraj at **antony.jeyaraj@dsv.com**. He can confirm available space and assist with pricing or scheduling a visit."})
# --- Warehouse Temperature Zones ---
    if match([
    r"\btemp\b", r"temperture", r"temperature", r"temperature zones", r"warehouse temp", 
    r"warehouse temperature", r"cold room", r"freezer room", r"ambient temperature",
    r"temp.*zones", r"how cold", r"cold storage", r"temperature range"]):
        return jsonify({"reply":
        "DSV warehouses support three temperature zones:\n\n"
        "🟢 **Ambient Storage**: +18°C to +25°C — for general cargo and FMCG\n"
        "🔵 **Cold Room**: +2°C to +8°C — for food and pharmaceuticals\n"
        "🔴 **Freezer Room**: –22°C — for frozen goods and sensitive materials\n\n"
        "All temperature-controlled areas are monitored 24/7 and GDP-compliant."})

    if match([r"\btapa\b", r"tapa certified", r"tapa standard", r"tapa compliance"]):
        return jsonify({"reply": "TAPA stands for Transported Asset Protection Association. It’s a global security standard for the safe handling, warehousing, and transportation of high-value goods. DSV follows TAPA-aligned practices for secure transport and facility operations, including access control, CCTV, sealed trailer loading, and secured parking."})
    if match([r"freezone", r"free zone", r"abu dhabi freezone", r"airport freezone", r"freezone warehouse"]):
        return jsonify({"reply": "DSV operates a GDP-compliant warehouse in the **Abu Dhabi Airport Freezone**, specialized in pharmaceutical and healthcare logistics. It offers:\n- Temperature-controlled and cold chain storage\n- Customs-cleared import/export operations\n- Proximity to air cargo terminals\n- Full WMS and track-and-trace integration\nThis setup supports fast, regulated distribution across the UAE and GCC."})
    if match([r"\bqhse\b", r"quality health safety environment", r"qhse policy", r"qhse standards", r"dsv qhse"]):
        return jsonify({"reply": "DSV follows strict QHSE standards across all operations. This includes:\n- Quality checks (ISO 9001)\n- Health & safety compliance (ISO 45001)\n- Environmental management (ISO 14001)\nAll staff undergo QHSE training, and warehouses are equipped with emergency protocols, access control, firefighting systems, and first-aid kits."})
    if match([r"\bhse\b", r"health safety environment", r"dsv hse", r"hse policy", r"hse training"]):
        return jsonify({"reply": "DSV places strong emphasis on HSE compliance. We implement:\n- Safety inductions and daily toolbox talks\n- Fire drills and emergency response training\n- PPE usage and incident reporting procedures\n- Certified HSE officers across sites\nWe’re committed to zero harm in the workplace."})
    if match([r"training", r"staff training", r"employee training", r"warehouse training", r"qhse training"]):
        return jsonify({"reply": "All DSV warehouse and transport staff undergo structured training programs, including:\n- QHSE training (Safety, Fire, First Aid)\n- Equipment handling (Forklifts, Cranes, VNA)\n- WMS and inventory systems\n- Customer service and operational SOPs\nRegular refresher courses are also conducted."})
    if match([r"\bdg\b", r"dangerous goods", r"hazardous material", r"hazmat", r"hazard class", r"dg storage"]):
        return jsonify({"reply": "Yes, DSV handles **DG (Dangerous Goods)** and hazardous materials in specialized chemical storage areas. We comply with all safety and documentation requirements including:\n- Hazard classification and labeling\n- MSDS (Material Safety Data Sheet) submission\n- Trained staff for chemical handling\n- Temperature-controlled and fire-protected zones\n- Secure access and emergency systems\n\nPlease note: For a DG quotation, we require the **material name, hazard class, CBM, period, and MSDS**."})

    # --- Chamber Mapping ---
    if match([r"ch2|chamber 2"]):
        return jsonify({"reply": "Chamber 2 is used by PSN (Federal Authority of Protocol and Strategic Narrative)."})
    if match([r"ch3|chamber 3"]):
        return jsonify({"reply": "Chamber 3 is used by food clients and fast-moving items."})
    if match([r"who.*in.*chamber|who.*in.*ch\d+"]):
        return jsonify({"reply": "The chambers in 21K warehouse are:\nCh1 – Khalifa University\nCh2 – PSN\nCh3 – Food clients\nCh4 – MCC, TR, ADNOC\nCh5 – PSN\nCh6 – ZARA, TR\nCh7 – Civil Defense & RMS"})
        
    # --- Warehouse Occupancy ---
    if match([r"warehouse occupancy|occupancy|space available|any space in warehouse|availability.*storage"]):
        return jsonify({"reply": "For warehouse occupancy, contact Biju Krishnan at biju.krishnan@dsv.com."})
    if match([r"open yard.*occupancy|yard space.*available|yard capacity|yard.*availability"]):
        return jsonify({"reply": "For open yard occupancy, contact Antony Jeyaraj at antony.jeyaraj@dsv.com."})
# --- Industry: Retail & Fashion ---
    if match([r"\bretail\b", r"fashion and retail", r"fashion logistics", r"retail supply chain"]):
        return jsonify({"reply": "DSV provides tailored logistics solutions for the **retail and fashion industry**, including:\n- Warehousing (racked, ambient, VNA)\n- Inbound & outbound transport\n- Value Added Services (labeling, repacking, tagging)\n- Last-mile delivery to malls and retail stores\n- WMS integration for real-time visibility"})

# --- Industry: Oil & Gas ---
    if match([r"oil and gas", r"oil & gas", r"\bo&g\b", r"energy sector", r"oil logistics"]):
        return jsonify({"reply": "DSV supports the **Oil & Gas industry** across Abu Dhabi and the GCC through:\n- Storage of chemicals and DG\n- Heavy equipment transport\n- 3PL/4PL project logistics\n- ADNOC-compliant warehousing and safety\n- Support for offshore & EPC contractors with specialized fleet"})

# --- Industry: Breakbulk / Heavy Logistics ---
    if match([r"breakbulk", r"break bulk", r"heavy cargo", r"non-containerized cargo"]):
        return jsonify({"reply": "DSV handles **breakbulk and heavy logistics** including:\n- Oversized cargo (machinery, steel, transformers)\n- Lowbed trailer and crane support\n- Project logistics & site delivery\n- DG compliance and route planning\n- Full UAE & GCC transport coordination"})
    if match([r"last mile", r"last mile delivery", r"final mile", r"city delivery"]):
        return jsonify({"reply": "DSV offers **last-mile delivery** services across the UAE using small city trucks and vans. These are ideal for e-commerce, retail, and healthcare shipments requiring fast and secure delivery to final destinations. Deliveries are WMS-tracked and coordinated by our OCC team for full visibility."})
    if match([r"\binventory\b", r"inventory management", r"inventory control", r"inventory system", r"stock tracking"]):
        return jsonify({"reply": "DSV uses INFOR WMS to manage all inventory activities. It provides:\n- Real-time stock visibility\n- Bin-level tracking\n- Batch/serial number control\n- Expiry tracking (for pharma/FMCG)\n- Integration with your ERP system"})
    if match([r"cross dock", r"cross docking", r"cross-dock", r"crossdock facility"]):
        return jsonify({"reply": "Yes, DSV supports **cross-docking** for fast-moving cargo:\n- Receive → Sort → Dispatch (no storage)\n- Ideal for FMCG, e-commerce, and retail\n- Reduces lead time and handling\n- Available at Mussafah and KIZAD hubs"})
    if match([r"transit store", r"transit warehouse", r"transit storage", r"temporary storage", r"short term storage"]):
        return jsonify({"reply": "DSV offers **transit storage** for short-term cargo holding. Ideal for:\n- Customs-cleared goods awaiting dispatch\n- Re-export shipments\n- Short-duration contracts\nOptions available in Mussafah, Airport Freezone, and KIZAD."})

    # --- EV trucks ---
    if match([r"ev truck|electric vehicle|zero emission|sustainable transport"]):
        return jsonify({"reply": "DSV Abu Dhabi operates EV trucks hauling 40ft containers. Each has ~250–300 km range and supports port shuttles & green logistics."})

    # --- DSV Managing Director (MD) ---
    if match([r"\bmd\b|managing director|head of dsv|ceo|boss of dsv|hossam mahmoud"]):
        return jsonify({"reply": "Mr. Hossam Mahmoud is the Managing Director, Road & Solutions and CEO Abu Dhabi. He oversees all logistics, warehousing, and transport operations in the region."})

    # --- What is WMS ---
    if match([r"what is wms|wms meaning|warehouse management system"]):
        return jsonify({"reply": "WMS stands for Warehouse Management System. DSV uses INFOR WMS for inventory control, inbound/outbound, and full visibility."})

    # --- Services DSV Provides ---
    if match([
    r"what.*service[s]?.*dsv.*provide",
    r"what (do|does).*dsv.*do",
    r"what (do|does).*they.*do",
    r"what (do|does).*they.*serve",
    r"what (do|does).*they.*offer", 
    r"what.*service[s].*they.*provide",
    r"dsv.*offer", 
    r"dsv.*specialize", 
    r"dsv.*work", 
    r"dsv.*services", 
    r"what.*type.*service", 
    r"type.*of.*logistics", 
    r"services.*dsv", 
    r"what.*dsv.*do", 
    r"dsv.*offerings"]):
        return jsonify({"reply": 
        "**DSV Abu Dhabi** provides full logistics and supply chain services, including:\n\n"
        "🚚 **2PL** – Road transport, containers, last-mile delivery\n"
        "🏢 **3PL** – Warehousing, inventory, VAS, WMS\n"
        "🔗 **3.5PL** – Hybrid logistics (execution + partial strategy)\n"
        "🧠 **4PL** – Fully managed supply chain operations\n\n"
        "**Main Facilities:**\n"
        "- 📍 **21K Warehouse (Mussafah)** – 21,000 sqm, 7 chambers\n"
        "- 📍 **M44 / M45** – Sub-warehouses in Mussafah\n"
        "- 📍 **Al Markaz (Hameem)** – Regional support\n"
        "- 📍 **KIZAD** – 360,000 sqm open yard\n"
        "- 📍 **Airport Freezone** – GDP-compliant storage for healthcare\n\n"
        "📞 +971 2 555 2900 | 🌐 [dsv.com](https://www.dsv.com)"})

    # --- What does DSV mean ---
    if match([
    r"\bdsv\b(?!.*fleet)", r"about dsv", r"who is dsv", r"what is dsv", r"dsv info", r"dsv abu dhabi",
    r"tell me about dsv", r"dsv overview", r"dsv abbreviation", r"dsv stands for", r"what does dsv mean"]):
        return jsonify({"reply":
        "DSV stands for **'De Sammensluttede Vognmænd'**, meaning **'The Consolidated Hauliers'** in Danish. "
        "Founded in 1976, DSV is a global logistics leader operating in over 80 countries."})

# --- DSV Sustainability Vision ---
    if match([
    r"sustainability", r"green logistics", r"sustainable practices", r"environmental policy",
    r"carbon footprint", r"eco friendly", r"zero emission goal", r"environment commitment"]):
        return jsonify({"reply":
    "DSV is committed to **sustainability and reducing its environmental footprint** across all operations. Initiatives include:\n"
    "- Transition to **electric vehicles (EV)** for last-mile and container transport\n"
    "- Use of **solar energy** and energy-efficient warehouse lighting\n"
    "- Consolidated shipments to reduce CO₂ emissions\n"
    "- Compliance with **ISO 14001** (Environmental Management)\n"
    "- Green initiatives in packaging, recycling, and process optimization\n\n"
    "DSV’s global strategy aligns with the UN Sustainable Development Goals and aims for net-zero emissions by 2050."})

    # --- Industry Tags (FMCG, Insurance, Healthcare, Ecommerce) ---
    if match([r"\bfmcg\b|fast moving|consumer goods"]):
        return jsonify({"reply": "DSV provides fast turnaround warehousing for FMCG clients including dedicated racking, SKU control, and high-frequency dispatch."})
    if match([r"insurance|is insurance included|cargo insurance"]):
        return jsonify({"reply": "Insurance is not included by default in quotations. It can be arranged separately upon request."})

    # --- Lean Six Sigma ---
    if match([r"lean six sigma|warehouse improvement|continuous improvement|kaizen|process efficiency|6 sigma|warehouse process improvement|lean method"]):
        return jsonify({"reply": "DSV applies Lean Six Sigma principles in warehouse design and process flow to reduce waste, improve accuracy, and maximize efficiency. We implement 5S, KPI dashboards, and root-cause analysis for continuous improvement."})

    # --- Warehouse Activities ---
    if match([r"warehouse activities|inbound process|outbound process|wh process|warehouse process|SOP|operation process|putaway|replenishment|picking|packing|cycle count"]):
        return jsonify({"reply": "Warehouse activities include:\n- Inbound: receiving, inspection, putaway\n- Outbound: picking, packing, dispatch\n- Replenishment, cycle counting, returns, VAS, and system updates via WMS."})
    if match([r"warehouse temp|temperature.*zone|storage temperature|cold room|freezer|ambient temp|warehouse temperature"]):
        return jsonify({"reply": "DSV provides 3 temperature zones:\n- **Ambient**: +18°C to +25°C\n- **Cold Room**: +2°C to +8°C\n- **Freezer**: –22°C\nThese zones are used for FMCG, pharmaceuticals, and temperature-sensitive products."})
    if match([r"size of our warehouse|total warehouse area|total sqm|warehouse size|how big.*warehouse"]):
        return jsonify({"reply": "DSV Abu Dhabi has approx. **44,000 sqm** of warehouse space:\n- 21K in Mussafah (21,000 sqm)\n- M44 (5,760 sqm)\n- M45 (5,000 sqm)\n- Al Markaz in Hameem (12,000 sqm)\nPlus 360,000 sqm of open yard."})
    if match([r"\bwh process\b", r"warehouse process", r"warehouse operations", r"warehouse workflow", r"\bwh\b.*operation", r"warehouse tasks", r"warehouse flow"]):
        return jsonify({"reply": "Typical warehouse processes at DSV include:\n1️⃣ **Inbound**: receiving, inspection, put-away\n2️⃣ **Storage**: in racks or bulk zones\n3️⃣ **Order Processing**: picking, packing, labeling\n4️⃣ **Outbound**: staging, dispatch, transport coordination\n5️⃣ **Inventory Control**: cycle counting, stock checks, and returns\n\nAll activities are managed through our INFOR WMS system for full visibility and traceability."})
    if match([r"kitting", r"assembly", r"kitting and assembly", r"value added kitting"]):
        return jsonify({"reply": "DSV provides **kitting and assembly** as a Value Added Service:\n- Combine multiple SKUs into kits\n- Light assembly of components\n- Repacking and labeling\n- Ideal for retail, pharma, and project logistics"})
    if match([r"packing material", r"what packing material", r"materials used for packing"]):
        return jsonify({"reply": "DSV uses high-grade packing materials:\n- Shrink wrap (6 rolls per box, 1 roll = 20 pallets)\n- Strapping rolls + buckle kits (1 roll = 20 pallets)\n- Bubble wrap, carton boxes, foam sheets\n- Heavy-duty pallets (wooden/plastic)\nUsed for relocation, storage, and export."})
    if match([r"\brelocation\b", r"move warehouse", r"shift cargo", r"site relocation"]):
        return jsonify({"reply": "Yes, DSV provides full **relocation services**:\n- Machinery shifting\n- Office and warehouse relocations\n- Packing, transport, offloading\n- Insurance and dismantling available\nHandled by our trained team with all safety measures."})

    # --- Machinery / Machineries ---
    if match([r"machinery|machineries|machines used|equipment|equipment used"]):
        return jsonify({"reply": "DSV uses forklifts (3–15T), VNA, reach trucks, pallet jacks, cranes, and container lifters in warehouse and yard operations."})

    if match([r"pallet.*bay|how many.*bay.*pallet", r"bay.*standard pallet", r"bay.*euro pallet"]):
        return jsonify({"reply": "Each bay in 21K can accommodate 14 Standard pallets or 21 Euro pallets. This layout maximizes efficiency for various cargo sizes."})

    # --- DSV Ecommerce, Insurance, WMS ---
    if match([r"ecommerce|e-commerce|online retail|ecom|dsv online|shop logistics|online order|fulfillment center"]):
        return jsonify({"reply": "DSV provides end-to-end e-commerce logistics including warehousing, order fulfillment, pick & pack, returns handling, last-mile delivery, and integration with Shopify, Magento, and custom APIs. Our Autostore and WMS systems enable fast, accurate processing of online orders from our UAE hubs including KIZAD and Airport Freezone."})

    if match([r"insurance|cargo insurance|storage insurance|are items insured"]):
        return jsonify({"reply": "Insurance is not included by default in DSV storage or transport quotes. It can be arranged upon client request, and is subject to cargo value, category, and terms agreed."})

    if match([r"\bwms\b|warehouse management system|inventory software|tracking system|dsv.*system"]):
        return jsonify({"reply": "DSV uses the INFOR Warehouse Management System (WMS) to manage inventory, inbound/outbound flows, and order tracking. It supports real-time dashboards and client integration."})

    if match([r"warehouse activities|warehouse tasks|daily warehouse work"]):
        return jsonify({"reply": "DSV warehouse activities include receiving (inbound), put-away, storage, replenishment, order picking, packing, staging, and outbound dispatch. We also handle inventory audits, cycle counts, and VAS."})

    if match([r"warehouse process|inbound|outbound|putaway|replenishment|dispatch"]):
        return jsonify({"reply": "Typical warehouse processes at DSV: (1) Inbound receiving, (2) Put-away into racks or zones, (3) Order picking or replenishment, (4) Packing & labeling, (5) Outbound dispatch. All steps are WMS-tracked."})
    if match([r"\bsop\b", r"standard operating procedure", r"standard operation process"]):
        return jsonify({"reply": 
        "SOP stands for **Standard Operating Procedure**. It refers to detailed, written instructions to achieve uniformity in operations. "
        "DSV maintains SOPs for all warehouse, transport, and VAS processes to ensure safety, compliance, and efficiency."})

# --- Air & Sea Services ---
    if match([
    r"air and sea", r"sea and air", r"air & sea", r"air freight and sea freight",
    r"dsv air and sea", r"dsv sea and air", r"dsv air & sea", r"air ocean", r"air & ocean"]):
        return jsonify({"reply": 
        "DSV provides comprehensive **Air & Sea freight forwarding** services globally and in the UAE:\n\n"
        "✈️ **Air Freight**:\n"
        "- Express, standard, and consolidated options\n"
        "- Charter solutions for urgent cargo\n"
        "- Abu Dhabi Airport Freezone warehouse integration\n\n"
        "🚢 **Sea Freight**:\n"
        "- Full Container Load (FCL) and Less than Container Load (LCL)\n"
        "- Customs clearance and documentation support\n"
        "- Direct access to UAE ports via Jebel Ali, Khalifa, and Zayed Port\n\n"
        "Our team handles end-to-end transport, consolidation, and global forwarding through DSV’s global network."})

    # --- chemical quotation ---
    if match([
    r"what.*(need|have).*collect.*chemical.*quote",
    r"what.*(to|do).*collect.*chemical.*quotation",
    r"build.*up.*chemical.*quote", r"build.*chemical.*quote",
    r"make.*chemical.*quotation", r"prepare.*chemical.*quote",
    r"chemical.*quote.*requirements", r"requirements.*chemical.*quote",
    r"info.*for.*chemical.*quote", r"details.*for.*chemical.*quotation",
    r"what.*required.*chemical.*quotation", r"quotation.*chemical.*details"]):
        return jsonify({"reply":
        "To provide a quotation for **chemical storage**, please collect the following from the client:\n"
        "1️⃣ **Product Name & Type**\n"
        "2️⃣ **Hazard Class / Classification**\n"
        "3️⃣ **Required Volume (CBM/SQM)**\n"
        "4️⃣ **Storage Duration (contract period)**\n"
        "5️⃣ **MSDS** – Material Safety Data Sheet\n"
        "6️⃣ **Any special handling or packaging needs**"})
    if match([r"store.*chemical|quotation.*chemical|data.*chemical|requirement.*chemical"]):
        return jsonify({"reply": "To quote for chemical storage, we need:\n- Material name\n- Hazard class\n- CBM\n- Period\n- MSDS (Material Safety Data Sheet)."})
    if match([r"\bmsds\b|material safety data sheet|chemical data"]):
        return jsonify({"reply": "Yes, MSDS (Material Safety Data Sheet) is mandatory for any chemical storage inquiry. It ensures safe handling and classification of the materials stored in DSV’s facilities."})
    if match([r"quote.*chemical.*warehouse", r"quote.*chemical storage", r"quote.*any storage", r"what.*need.*quote.*storage", r"build.*quote.*chemical"]):
        return jsonify({"reply":
        "To build a quotation for storage (especially chemical), collect the following:\n"
        "1️⃣ Type of material / hazard class\n"
        "2️⃣ Volume (CBM or SQM)\n"
        "3️⃣ Storage duration (contract period)\n"
        "4️⃣ MSDS if chemical\n"
        "5️⃣ Handling frequency (thruput)\n\n"
        "Once ready, please fill the form on the left."})

# --- General 3PL Quotation Requirement ---
    if match([
    r"(what.*collect.*client.*quotation)", r"(what.*info.*client.*quote)", 
    r"(quotation.*requirements)", r"(quotation.*information.*client)", 
    r"(details.*for.*quotation)", r"(build.*quotation.*info)", 
    r"(prepare.*quotation.*client)", r"(required.*info.*quote)"]):
        return jsonify({"reply": 
        "To build a proper 3PL storage quotation, please collect the following information from the client:\n"
        "1️⃣ **Type of Commodity** – What items are being stored (FMCG, chemical, pharma, etc.)\n"
        "2️⃣ **Contract Period** – Expected duration of the agreement (in months or years)\n"
        "3️⃣ **Storage Volume** – In CBM/day, CBM/month, or CBM/year for warehousing; in SQM for open yard\n"
        "4️⃣ **Throughput Volumes (IN/OUT)** – Daily or monthly volume in CBM to determine handling pattern and frequency\n\n"
        "Once these details are available, you can proceed to fill the main form to generate a quotation."})
    if match([r"proposal|quotation|quote.*open yard|send me.*quote|how to get quote|need.*quotation"]):
        return jsonify({"reply": "To get a full quotation, please close this chat and fill the details in the main form on the left. The system will generate a downloadable document for you."})
 # --- service ---
    if match([r"\bwhat is 2pl\b", r"\b2pl\b", r"second party logistics", r"what is 2pl", r"2pl meaning"]):
        return jsonify({"reply": "2PL (Second Party Logistics) refers to transport-only logistics providers. DSV’s 2PL services include dedicated trucking, container movement, and distribution across the UAE and GCC using our own fleet."})
    if match([r"\bwhat is 3pl\b", r"\b3pl\b", r"third party logistics"]):
        return jsonify({"reply": "3PL (Third Party Logistics) involves outsourcing logistics operations such as warehousing, transportation, picking/packing, and order fulfillment to a provider like DSV."})
    if match([r"\bwhat is 4pl\b", r"\b4pl\b", r"fourth party logistics"]):
        return jsonify({"reply": "4PL (Fourth Party Logistics) is a fully integrated supply chain solution where DSV manages all logistics operations, partners, systems, and strategy on behalf of the client. DSV acts as a single point of contact and coordination."})
    if match([r"\bwhat is 3.5pl\b", r"\b3.5pl\b", r"three and half pl", r"3pl plus", r"middle of 3pl and 4pl"]):
        return jsonify({"reply": "3.5PL is an emerging term referring to a hybrid between **3PL and 4PL**:\n- DSV provides operational execution like a 3PL\n- And partial strategic control like a 4PL\nIdeal for clients wanting control with partial outsourcing."})

    # --- Transportation---
    if match([
    r"\bfleet\b", r"dsv.*fleet", r"fleet.*dsv",
    r"\bdsv transportation\b", r"truck fleet", r"transport fleet", r"fleet info",
    r"what.*fleet.*dsv", r"dsv.*trucks", r"types of fleet"]):
        return jsonify({"reply": 
        "DSV operates a large fleet in the UAE including:\n\n"
        "- 🚛 Flatbed trailers\n"
        "- 📦 Box trucks\n"
        "- 🚚 Double trailers\n"
        "- ❄️ Refrigerated trucks (chiller/freezer)\n"
        "- 🏗 Lowbeds\n"
        "- 🪨 Tippers\n"
        "- 🏙 Small city delivery trucks\n\n"
        "Fleet vehicles support all types of transport including full truckload (FTL), LTL, and container movements."})

    if match([r"\bfleet\b", r"\bdsv fleet\b",r"\bdsv transportation\b", r"truck fleet", r"transport fleet", r"fleet info"]):
        return jsonify({"reply": "DSV operates a large fleet in the UAE including:\n- Flatbed trailers\n- Box trucks\n- Double trailers\n- Refrigerated trucks (chiller/freezer)\n- Lowbeds\n- Tippers\n- Small city delivery trucks\nFleet vehicles support all types of transport including full truckload (FTL), LTL, and container movements."})
    if match([r"truck types", r"trucks", r"transportation types", r"dsv trucks", r"transport.*available", r"types of transport", r"trucking services"]):
        return jsonify({"reply": "DSV provides local and GCC transportation using:\n- Flatbeds for general cargo\n- Lowbeds for heavy equipment\n- Tippers for construction bulk\n- Box trucks for secure goods\n- Refrigerated trucks for temperature-sensitive cargo\n- Double trailers for long-haul\n- Vans and city trucks for last-mile delivery."})
    if match([r"\btransportation\b", r"tell me about transportation", r"transport services", r"what is transportation", r"dsv transportation"]):
        return jsonify({"reply":
        "DSV offers full-service land transportation across the UAE and GCC. We operate a modern fleet including:\n"
        "- 🚛 Flatbeds (up to 25 tons)\n"
        "- 🏗 Lowbeds for heavy or oversized cargo\n"
        "- 🪨 Tippers for bulk material (sand, gravel, etc.)\n"
        "- 📦 Box trucks for protected cargo\n"
        "- ❄️ Reefer trucks for temperature-controlled delivery\n"
        "- 🚚 Double trailers for high-volume long-haul moves\n"
        "- 🏙 Small city trucks for last-mile distribution\n\n"
        "All transport is coordinated by our OCC team in Abu Dhabi with real-time tracking, WMS integration, and documentation support."})

    # --- UAE Emirates Distance + Travel Time (Individual Matches) ---
    if match([r"abu dhabi.*dubai|dubai.*abu dhabi"]):
        return jsonify({"reply": "The distance between Abu Dhabi and Dubai is about **140 km**, and the travel time is approximately **1.5 hours**."})
    if match([r"abu dhabi.*sharjah|sharjah.*abu dhabi"]):
        return jsonify({"reply": "The distance between Abu Dhabi and Sharjah is about **160 km**, and the travel time is approximately **1.5 to 2 hours**."})
    if match([r"abu dhabi.*ajman|ajman.*abu dhabi"]):
        return jsonify({"reply": "The distance between Abu Dhabi and Ajman is approximately **170 km**, with a travel time of about **1.5 to 2 hours**."})
    if match([r"abu dhabi.*ras al khaimah|ras al khaimah.*abu dhabi|rak.*abu dhabi|abu dhabi.*rak"]):
        return jsonify({"reply": "The road distance from Abu Dhabi to Ras Al Khaimah is about **240 km**, and the travel time is around **2.5 to 3 hours**."})
    if match([r"abu dhabi.*fujairah|fujairah.*abu dhabi"]):
        return jsonify({"reply": "Abu Dhabi to Fujairah is approximately **250 km**, with a travel time of about **2.5 to 3 hours**."})
    if match([r"dubai.*sharjah|sharjah.*dubai"]):
        return jsonify({"reply": "Dubai to Sharjah is around **30 km**, and the travel time is typically **30 to 45 minutes**."})
    if match([r"dubai.*ajman|ajman.*dubai"]):
        return jsonify({"reply": "Dubai to Ajman is approximately **40 km**, and it takes around **40 to 50 minutes** by road."})
    if match([r"dubai.*ras al khaimah|ras al khaimah.*dubai|dubai.*rak|rak.*dubai"]):
        return jsonify({"reply": "The distance between Dubai and Ras Al Khaimah is around **120 km**, with a travel time of **1.5 to 2 hours**."})
    if match([r"dubai.*fujairah|fujairah.*dubai"]):
        return jsonify({"reply": "Dubai to Fujairah is approximately **130 km**, and the travel time is about **2 hours**."})
    if match([r"sharjah.*ajman|ajman.*sharjah"]):
        return jsonify({"reply": "Sharjah and Ajman are extremely close — only about **15 km**, with a travel time of **15 to 20 minutes**."})
    if match([r"sharjah.*fujairah|fujairah.*sharjah"]):
        return jsonify({"reply": "Sharjah to Fujairah is roughly **110 km**, and takes about **2 hours** by road."})
    if match([r"sharjah.*ras al khaimah|ras al khaimah.*sharjah|sharjah.*rak|rak.*sharjah"]):
        return jsonify({"reply": "Sharjah to Ras Al Khaimah is approximately **100 km**, and the travel time is around **1.5 to 2 hours**."})

    if match([
    r"truck capacity", r"how many ton", r"truck tonnage", r"truck.*can carry", r"truck load",
    r"flatbed.*ton", r"flatbed.*load", r"flatbed capacity",
    r"double trailer.*ton", r"articulated.*capacity",
    r"box truck.*ton", r"curtainside.*load", r"box truck capacity",
    r"reefer.*ton", r"refrigerated truck.*capacity", r"chiller truck.*load",
    r"city truck.*ton", r"1 ton truck", r"3 ton truck",
    r"lowbed.*ton", r"lowbed.*capacity",
    r"tipper.*ton", r"dump truck.*load", r"bulk truck.*ton"]):
        return jsonify({"reply": "Here’s the typical tonnage each DSV truck type can carry:\n\n"
        "🚛 **Flatbed Truck**: up to 22–25 tons (ideal for general cargo, pallets, containers)\n"
        "🚚 **Double Trailer (Articulated)**: up to 50–60 tons combined (used for long-haul or inter-emirate)\n"
        "📦 **Box Truck / Curtainside**: ~5–10 tons (weather-protected for packaged goods)\n"
        "❄️ **Refrigerated Truck (Reefer)**: 3–12 tons depending on size (for temperature-sensitive cargo)\n"
        "🏙 **City Truck (1–3 Ton)**: 1 to 3 tons (last-mile delivery within cities)\n"
        "🏗 **Lowbed Trailer**: up to 60 tons (for heavy equipment and oversized machinery)\n"
        "🪨 **Tipper / Dump Truck**: ~15–20 tons (for bulk cargo like sand, gravel, or construction material)"})

    if match([r"(distance|how far|km).*mussafah.*(al markaz|markaz|hameem|hamim|ghayathi|ruwais|mirfa|madinat zayed|western region)"]):
        return jsonify({"reply": (
        "Approximate road distances from Mussafah:\n"
        "- Al Markaz: **60 km**\n"
        "- Hameem: **90 km**\n"
        "- Madinat Zayed: **150 km**\n"
        "- Mirfa: **140 km**\n"
        "- Ghayathi: **240 km**\n"
        "- Ruwais: **250 km**\n"
        "\nLet me know if you need travel time or transport support too.")})
        # --- DSV Abu Dhabi Facility Sizes ---
    if match([
        r"plot size", r"abu dhabi total area", r"site size", r"facility size", r"total sqm", r"how big",
        r"yard size", r"open yard area", r"size of open yard", r"open yard.*size", r"area of open yard"]):
        return jsonify({"reply": "DSV Abu Dhabi's open yard spans 360,000 SQM across Mussafah and KIZAD. The total logistics plot is 481,000 SQM, including 100,000 SQM of service roads and utilities, and a 21,000 SQM warehouse (21K)."})

    if match([r"sub warehouse|m44|m45|al markaz|abu dhabi warehouse total|all warehouses"]):
        return jsonify({"reply": "In addition to the main 21K warehouse, DSV operates sub-warehouses in Abu Dhabi: M44 (5,760 sqm), M45 (5,000 sqm), and Al Markaz (12,000 sqm). Combined with 21K, the total covered warehouse area in Abu Dhabi is approximately 44,000 sqm."})

    if match([r"terms and conditions|quotation policy|T&C|billing cycle|operation timing|payment terms|quotation validity"]):
        return jsonify({"reply": "DSV quotations include the following terms: Monthly billing, final settlement before vacating, 15-day quotation validity, subject to space availability. The depot operates Monday–Friday 8:30 AM to 5:30 PM. Insurance is not included by default. An environmental fee of 0.15% is added to all invoices. Non-moving cargo over 3 months may incur extra storage tariff."})
    # --- QHSE ---   
    if match([r"safety training|warehouse training|fire drill|manual handling|staff safety|employee training|toolbox talk"]):
        return jsonify({"reply": "DSV staff undergo regular training in fire safety, first aid, manual handling, emergency response, and site induction. We also conduct toolbox talks and refresher sessions to maintain safety awareness and operational excellence."})
    # --- DSV & ADNOC Relationship ---
    if match([r"adnoc|adnoc project|dsv.*adnoc|oil and gas project|dsv support.*adnoc|logistics for adnoc"]):
        return jsonify({"reply": "DSV has an active relationship with ADNOC and its group companies, supporting logistics for Oil & Gas projects across Abu Dhabi. This includes warehousing of chemicals, fleet transport to remote sites, 3PL for EPC contractors, and marine logistics for ADNOC ISLP and offshore projects. All operations are QHSE compliant and meet ADNOC’s safety and performance standards."})
    # --- UAE Summer Midday Break ---
    if match([r"summer break|midday break|working hours summer|12.*3.*break|uae heat ban|no work afternoon|hot season schedule"]):
        return jsonify({"reply": "DSV complies with UAE summer working hour restrictions. From June 15 to September 15, all outdoor work (including open yard and transport loading) is paused daily between 12:30 PM and 3:30 PM. This ensures staff safety and follows MOHRE guidelines."})
    # --- Client Name Queries ---
    if match([r"chambers.*21k", r"how many.*chambers", r"warehouse.*layout", r"wh.*layout", r"warehouse.*structure", r"wh.*layout", r"\bchambers\b"]):
        return jsonify({"reply": "There are 7 chambers in the 21K warehouse with different sizes and rack type, chambers'size start from 1000 sqm upto 5000 sqm which can accomodate upto to total 35000 cbm"})

    if match([r"who is in ch(\d+)|who is in chamber(\d+)|21K clients|warehouse clients|WH clients|client in ch(\d+)|ch\d+"]):
        ch_num = re.search(r"ch(\d+)", message)
        if ch_num:
            chamber = int(ch_num.group(1))
            clients = {
                1: "Khalifa University",
                2: "PSN",
                3: "Food clients & fast-moving items",
                4: "MCC, TR, and ADNOC",
                5: "PSN",
                6: "ZARA & TR",
                7: "Civil Defense and the RMS"
            }
            client_name = clients.get(chamber, "unknown")
            return jsonify({"reply": f"Chamber {chamber} is occupied by {client_name}."})

    # --- Friendly Chat ---
    if match([r"\bhello\b|\bhi\b|\bhey\b|good morning|good evening"]):
        return jsonify({"reply": "Hello! I'm here to help with anything related to DSV logistics, transport, or warehousing."})
    if match([r"how.?are.?you|how.?s.?it.?going|whats.?up"]):
        return jsonify({"reply": "I'm doing great! How can I assist you with DSV services today?"})
    if match([r"\bthank(s| you)?\b|thx|appreciate"]):
        return jsonify({"reply": "You're very welcome! 😊"})

    # --- Fallback (never ask to rephrase) ---
    return jsonify({"reply": "I'm trained on everything related to DSV storage, transport, VAS, Mussafah warehouse, and services. Can you try asking again with more detail?"})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
